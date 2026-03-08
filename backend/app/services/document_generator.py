from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Iterable

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from app.core.config import settings
from app.services.file_utils import ensure_dir
from app.services.models import DocumentBundle, KnowledgeDocument


def _timestamp_link(timestamp: float, source_url: str | None) -> str:
    if not source_url:
        return f"{timestamp:.0f}s"
    sep = "&" if "?" in source_url else "?"
    return f"[{timestamp:.0f}s]({source_url}{sep}t={int(timestamp)})"


def _slugify(text: str) -> str:
    t = unicodedata.normalize("NFKD", text).lower()
    t = re.sub(r"[^\w\s-]", "", t, flags=re.UNICODE)
    t = re.sub(r"[-\s]+", "-", t).strip("-")
    return t or "section"


def render_markdown(doc: KnowledgeDocument) -> str:
    lines = [f"# {doc.title}", "", "## Overview", doc.overview, "", "## Table of Contents"]
    for c in doc.chapters:
        lines.append(f"- [{c.title}](#{_slugify(c.title)})")
    lines.append("")
    for c in doc.chapters:
        lines.append(f"## {c.title}")
        lines.append(f"Timestamp: {_timestamp_link(c.timestamp, c.source_url)}")
        lines.append("")
        lines.append(c.content)
        lines.append("")
    lines.append("## Summary")
    lines.append(doc.summary)
    lines.append("")
    return "\n".join(lines)


def generate_documents(doc: KnowledgeDocument, job_id: str) -> DocumentBundle:
    ensure_dir(settings.documents_dir)

    md = settings.documents_dir / f"{job_id}.md"
    content = render_markdown(doc)
    md.write_text(content, encoding="utf-8")

    docx = _generate_docx(doc, job_id, content)
    pdf = _generate_pdf_reportlab(job_id, content)

    return DocumentBundle(markdown_path=md, pdf_path=pdf, docx_path=docx)


def _generate_docx(doc: KnowledgeDocument, job_id: str, markdown_content: str) -> Path:
    p = settings.documents_dir / f"{job_id}.docx"
    d = Document()
    d.add_heading(doc.title, level=0)
    d.add_paragraph(doc.overview)
    d.add_paragraph("\n".join(markdown_content.splitlines()[4:]))
    d.save(p)
    return p


def _generate_pdf_reportlab(job_id: str, content: str) -> Path | None:
    pdf_path = settings.documents_dir / f"{job_id}.pdf"
    try:
        # Try Unicode fonts
        font_candidates = [
            Path("C:/Windows/Fonts/Nirmala.ttf"),
            Path("C:/Windows/Fonts/arialuni.ttf"),
            Path("C:/Windows/Fonts/segoeui.ttf"),
        ]
        font_file = next((f for f in font_candidates if f.exists()), None)
        if not font_file:
            return None

        pdfmetrics.registerFont(TTFont("UnicodeFont", str(font_file)))

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        x = 40
        y = height - 40
        c.setFont("UnicodeFont", 11)

        for raw in content.splitlines():
            line = raw if raw else " "
            # soft wrap
            while len(line) > 120:
                c.drawString(x, y, line[:120])
                y -= 16
                line = line[120:]
                if y < 40:
                    c.showPage()
                    c.setFont("UnicodeFont", 11)
                    y = height - 40
            c.drawString(x, y, line)
            y -= 16
            if y < 40:
                c.showPage()
                c.setFont("UnicodeFont", 11)
                y = height - 40

        c.save()
        return pdf_path
    except Exception:
        return None